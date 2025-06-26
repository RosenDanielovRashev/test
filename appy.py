import streamlit as st
from docx import Document
from io import BytesIO

# Session state: готов ли е документа
if "generated" not in st.session_state:
    st.session_state.generated = False

# Layout: две колони
col1, col2 = st.columns([1, 2])

# ЛЯВА КОЛОНА – въвеждане
with col1:
    st.header("Форма")
    name = st.text_input("Име")
    address = st.text_area("Адрес")

    if st.button("Генерирай документ"):
        if name.strip() != "" and address.strip() != "":
            st.session_state.generated = True
            st.session_state.name = name
            st.session_state.address = address
        else:
            st.warning("Моля, попълни всички полета!")

# ДЯСНА КОЛОНА – визуализация и сваляне
with col2:
    st.header("Документ")
    if st.session_state.generated:
        st.subheader("Преглед:")
        st.markdown(f"""
        #### Лични данни  
        **Име:** {st.session_state.name}  
        **Адрес:** {st.session_state.address}
        """)
        
        # Функция за създаване на DOCX
        def create_docx(name, address):
            doc = Document()
            doc.add_heading("Документ", 0)
            doc.add_paragraph(f"Име: {name}")
            doc.add_paragraph(f"Адрес: {address}")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        docx_file = create_docx(st.session_state.name, st.session_state.address)

        st.download_button(
            label="⬇️ Свали като DOCX",
            data=docx_file,
            file_name="личен_документ.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("Моля, въведи данни вляво и натисни 'Генерирай документ'")
